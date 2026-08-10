"""
utils.py
Small, dependency-light helper functions used across the app:
- saving uploads safely
- validating file types/sizes
- cleaning + validating the JSON that comes back from Gemini
- extracting text from .docx files (Gemini doesn't accept docx directly)
"""

import json
import re
import uuid
from pathlib import Path
from typing import Optional

from config import ALLOWED_EXTENSIONS, MAX_FILE_SIZE_MB, UPLOAD_DIR


class UnsupportedFileError(Exception):
    pass


class FileTooLargeError(Exception):
    pass


class InvalidModelResponseError(Exception):
    pass


def validate_upload(filename: str, size_bytes: int) -> None:
    """Raises a descriptive error if the uploaded file fails basic checks."""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise UnsupportedFileError(
            f"'{ext}' is not supported. Allowed types: "
            f"{', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    max_bytes = MAX_FILE_SIZE_MB * 1024 * 1024
    if size_bytes > max_bytes:
        raise FileTooLargeError(
            f"File is {size_bytes / (1024 * 1024):.1f} MB; "
            f"the limit is {MAX_FILE_SIZE_MB} MB."
        )


def save_uploaded_file(file_bytes: bytes, original_filename: str) -> Path:
    """Saves an uploaded file under a unique name and returns its path."""
    ext = Path(original_filename).suffix.lower()
    safe_name = f"{uuid.uuid4().hex}{ext}"
    dest = UPLOAD_DIR / safe_name
    dest.write_bytes(file_bytes)
    return dest


def extract_text_from_docx(path: Path) -> str:
    """Extracts plain text (paragraphs + tables) from a .docx file."""
    import docx  # python-docx

    document = docx.Document(str(path))
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def strip_markdown_fences(text: str) -> str:
    """
    Gemini is instructed to return raw JSON, but models occasionally wrap
    output in ```json ... ``` anyway. This removes that wrapping safely.
    """
    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def parse_and_validate_json(raw_text: str) -> dict:
    """
    Parses the model's raw text response into a dict and checks that the
    minimum required keys are present. Raises InvalidModelResponseError
    with a helpful message if parsing/validation fails.
    """
    cleaned = strip_markdown_fences(raw_text)

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise InvalidModelResponseError(
            f"Model response was not valid JSON: {exc}\n---\n{raw_text[:500]}"
        ) from exc

    required_top_level = {"overall_score", "grade", "questions", "summary"}
    missing = required_top_level - data.keys()
    if missing:
        raise InvalidModelResponseError(
            f"Model JSON is missing required keys: {sorted(missing)}"
        )

    if not isinstance(data["questions"], list) or len(data["questions"]) == 0:
        raise InvalidModelResponseError(
            "Model JSON contains no detected questions ('questions' is empty)."
        )

    required_question_keys = {
        "question_number", "question", "student_answer", "expected_solution",
        "analysis", "mistakes", "error_categories", "correct_answer",
        "suggestions", "difficulty", "score",
    }
    for i, q in enumerate(data["questions"]):
        missing_q = required_question_keys - q.keys()
        if missing_q:
            raise InvalidModelResponseError(
                f"Question #{i + 1} in model JSON is missing keys: {sorted(missing_q)}"
            )

    return data


def guess_mime_type(path: Path) -> Optional[str]:
    ext = path.suffix.lower()
    return {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".mp3": "audio/mpeg",
        ".wav": "audio/wav",
        ".m4a": "audio/mp4",
        ".ogg": "audio/ogg",
    }.get(ext)
